import vk_api
import requests
from bs4 import BeautifulSoup
from urllib.request import urlretrieve
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
import time
from config import login, password, vk_user_id, walls_count, walls_offset, size_photos, dictionars


# Создание документа word
def createDocumentWord(data_post):
    document = Document()

    # Текст
    document.add_paragraph(data_post['text'])

    # Дата создания поста
    date = document.add_paragraph(data_post['date'])
    date_format = date.paragraph_format
    date_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # # Тест текст
    # run = document.add_paragraph('Test text').add_run()
    # font = run.font
    # font.name = 'Calibri'
    # font.size = Pt(12)

    # # Ссылка
    # if 'link' in data_post:
    #     link = document.add_paragraph(data_post['link'])
    #     link_format = link.paragraph_format
    #     link_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Фото
    for file_name_photo in data_post['photos']:
        document.add_picture(file_name_photo, width=Inches(5.0))
        os.remove(file_name_photo)

    document.save(data_post['file_name'] + '.docx')


# Авто суммирование переменной walls_offset
def wallsOffsetPlus(walls_offset, walls_count):
    walls_offset = walls_offset + walls_count
    with open('config.py', 'a') as file: file.write('\nwalls_offset=' + str(walls_offset))


# Скачать файл
def downloadFile(url, file_name):
    # urlretrieve(url, file_name)
    r = requests.get(url, stream=True)

    with open(file_name, 'bw') as file:
        for chunk in r.iter_content(4096):
            file.write(chunk)


# Получение приложенной ссылки
def downloadAttachmentLink(attachment):
    return attachment['link']


# Организация скачивания вложений (Фото, Документов)
def downloadAttachmentDoc(attachment, attachment_photos):
    type = attachment['type']

    if type == 'photo':
        url = attachment['photo']['photo_604']
        file_name = url.rsplit('/', 1)[1]
        attachment_photos.append(file_name)

    elif type == 'doc':
        url = attachment[type]['url']
        file_name = attachment[type]['title']

    downloadFile(url, file_name)

    return attachment_photos


# Скачивание медиа контента (пока не реализовано и не используется)
def downloadAttachmentMedia(attachment, vk_session):
    type = attachment['type']
    if type == 'video':

        video_owner_id = attachment['video']['owner_id']
        video_id = attachment['video']['id']

        request_data = str(video_owner_id) + '_' + str(video_id)

        with vk_api.VkRequestsPool(vk_session) as pool:
            video_data = pool.method('video.get', {
                'videos': request_data
            })

        video_url = str(video_data.result['items'][0]['player'])
        new_video_url = video_url.replace('embed', 'watch')
        downloadFile(new_video_url, 'asdasd.avi')

        # html = requests.get(new_video_url).text
        # soup = BeautifulSoup(html, 'lxml')
        # print(soup)
        # # video_base_url = soup.find('div', id='sas')

    if type == 'audio':
        pass


# Определение категории через интерфейс коммандой строки
def identifyUserCategory(title, dictionars):
    set_category = input(title+'\n')
    for category in dictionars:
        try:
            category.index(set_category)
            return 'Статьи/' + category
        except ValueError:
            pass
    return 'Статьи/Скачанные'


# Имя файла
def createFileName(str):
    # Не должны заканчиваться точкой или пробелом
    simbols = '''☀$%b@&!-/\ "*?<>|:'''

    str = str[:50]

    for i in range(len(simbols)):
        str.replace(simbols[i], '')

    number_insert = str.find('http')

    if number_insert != -1:
        str = str[:number_insert]
        print(str)

    return str



# Логирование заголовка поста, где надо скачать медиа контент
def addPostInLog(data_post):
    text = data_post['title']
    date = data_post['date']

    data_insert = date + '\n' + text + '\n~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~\n'*2

    file = open('log_post_insert_media_content.txt', 'a')
    file.write(data_insert)
    file.close()


def main(dictionars):
    vk_session = vk_api.VkApi(login, password)

    try:
            vk_session.auth()
    except vk_api.AuthError as error_msg:
        print(error_msg)
        return

    with vk_api.VkRequestsPool(vk_session) as pool:
        user_walls = pool.method('wall.get', {
            'owner_id': vk_user_id,
            'count': walls_count,
            'offset': walls_offset
        })

    wallsOffsetPlus(walls_offset, walls_count)

    data_post = {}

    # Проход по всем постам
    for post in user_walls.result['items']:
        data_post['text'] = post['copy_history'][0]['text']
        data_post['title'] = data_post['text'][:200]
        data_post['file_name'] = createFileName(data_post['text'])
        data_post['date'] = time.strftime("%d-%m-%Y %H:%M:%S", time.localtime(post['copy_history'][0]['date']))

        # Получение всех типов вложений
        attachments_types = [attachment['type'] for attachment in post['copy_history'][0]['attachments']]

        # Логирование поста для скачавание медиа вложений
        if 'video' in attachments_types or 'audio' in attachments_types:
            os.chdir('Статьи/Скачанные')
            addPostInLog(data_post)

        # Смена каталога по категории
        dir = identifyUserCategory(data_post['title'], dictionars)
        print(dir+'\n~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~\n')
        data_post['dir'] = dir
        os.chdir(dir)


        # Создание и смена каталога при дополнительных вложениях
        if 'doc' in attachments_types or 'video' in attachments_types or 'audio' in attachments_types:
            os.mkdir(data_post['file_name'])
            os.chdir(data_post['file_name'])


        attachment_photos = []

        # Проход по всем вложениям поста
        for post_attachment in post['copy_history'][0]['attachments']:

            # Определение типа вложения
            type = post_attachment['type']

            if type == 'photo' or type == 'doc':
                attachment_photos = downloadAttachmentDoc(post_attachment, attachment_photos)

            elif type == 'link':
                data_post['link'] = downloadAttachmentLink(post_attachment)

            # elif type == 'audio' or type == 'video':
            #     downloadAttachmentMedia(post_attachment, vk_session)

        data_post['photos'] = attachment_photos

        # Создание документа MS Word
        createDocumentWord(data_post)


        # if os.path.exists(data_post['dir'] + data_post['file_name']):
        #     print('Ok')
        # else:
        #     print('No found file')


        # print(os.getcwd())
        # print(category)

        # import json
        # print(json.dumps(post_attachment, indent=2))


if __name__ == '__main__':
    main(dictionars)